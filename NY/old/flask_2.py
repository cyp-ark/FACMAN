from flask import Flask, request, jsonify, render_template, send_file
from influxdb_client import InfluxDBClient
from docx import Document
from docx.shared import Inches
from io import BytesIO
import base64
import openai

app = Flask(__name__)

# OpenAI 키와 InfluxDB 설정 입력
openai.api_key = "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"  

INFLUX_URL = "http://52.196.29.195:8086"
INFLUX_TOKEN = "CN2gilx5zoVNMPT5Qk8eSs_qlC6WFWX5mwSD0l0mw8o0k_FQKbL_g1aYu85HQxeaE_Ye44X_4EbIMu7ctpsKRQ=="
INFLUX_ORG = "org_kpmg"
client = InfluxDBClient(url=INFLUX_URL, token=INFLUX_TOKEN, org=INFLUX_ORG)

# 메인 페이지
@app.route("/")
def index():
    return render_template("test4.html")

# 보고서 생성 (LLM + DB 쿼리)
@app.route("/generate_report", methods=["POST"])
def generate_report():
    data = request.json
    process = data.get("process")
    range_str = data.get("range")

    query = f'''
    from(bucket: "{process}_status")
      |> range(start: -{range_str})
      |> filter(fn: (r) => r._measurement == "status_log")
      |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
      |> keep(columns: ["_time", "available", "event_type"])
    '''

    tables = client.query_api().query(query)

    total = 0
    available_sum = 0
    failure_count = 0
    time_labels = []
    available_values = []
    failure_values = []

    for table in tables:
        for record in table.records:
            available = record.values.get("available", 0)
            event_type = record.values.get("event_type", "")
            timestamp = record.values["_time"].strftime("%H:%M")

            total += 1
            available_sum += available
            if event_type == "failure":
                failure_count += 1

            time_labels.append(timestamp)
            available_values.append(round(available, 2))
            failure_values.append(1 if event_type == "failure" else 0)

    avg_avail = round((available_sum / total) * 100, 1) if total else 0

    prompt = f"""
공정명: {process}
기간: 최근 {range_str}
가동률 평균: {avg_avail}%
고장 횟수: {failure_count}회

위 데이터를 바탕으로 제조 공정 보고서를 작성해줘. 다음 항목을 포함해줘:
1. 공정 요약
2. 주요 이슈
3. 대응 조치
4. 향후 제언
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 제조공정 보고서를 작성하는 AI 비서야."},
                {"role": "user", "content": prompt}
            ]
        )
        return jsonify({
            "report": response.choices[0].message.content,
            "labels": time_labels,
            "available": available_values,
            "failures": failure_values
        })
    except Exception as e:
        print(f"오류 발생: {e}")
        return jsonify({"error": str(e)}), 500

# 보고서 .docx 파일 생성
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    try:
        data = request.json
        text = data.get("report", "")
        avail_img_b64 = data.get("availabilityImage", "")
        fail_img_b64 = data.get("failureImage", "")

        doc = Document()
        doc.add_heading("📄 스마트 제조 보고서", 0)
        doc.add_paragraph(text)

        def add_image(doc, b64_string, title):
            if b64_string and "base64," in b64_string:
                try:
                    doc.add_paragraph(title)
                    image_data = base64.b64decode(b64_string.split(",")[-1])
                    image_stream = BytesIO(image_data)
                    doc.add_picture(image_stream, width=Inches(2.75))
                except Exception as img_err:
                    print(f"이미지 디코딩 오류: {img_err}")

        add_image(doc, avail_img_b64, "✅ 가동률 변화")
        add_image(doc, fail_img_b64, "📊 고장 발생 분포")

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name="제조_보고서.docx")

    except Exception as e:
        print(f"📛 DOCX 생성 오류: {e}")
        return jsonify({"error": "파일 생성 실패"}), 500

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
