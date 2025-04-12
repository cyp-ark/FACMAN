import eventlet
eventlet.monkey_patch()

import os
import base64
from io import BytesIO
from flask import Flask, render_template, request, jsonify, send_file
from flask_socketio import SocketIO
from influxdb_client import InfluxDBClient
from openai import OpenAI
from dotenv import load_dotenv
from flask_cors import CORS
from docx import Document
from docx.shared import Inches
from datetime import datetime

from typing import Annotated
from typing_extensions import TypedDict
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.memory import MemorySaver
from langchain.agents import Tool
from langchain_openai import ChatOpenAI
from langchain_core.runnables import RunnableConfig

# ✅ 환경 변수 로드
load_dotenv()

# ✅ Flask 앱 및 SocketIO 초기화
app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode="eventlet")

# ✅ InfluxDB 클라이언트 설정
INFLUX_URL = os.getenv("INFLUX_URL")
INFLUX_TOKEN = os.getenv("INFLUX_TOKEN")
INFLUX_ORG = os.getenv("INFLUX_ORG")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

influx_client = InfluxDBClient(url=INFLUX_URL, token=INFLUX_TOKEN, org=INFLUX_ORG)
openai_client = OpenAI(api_key=OPENAI_API_KEY)

# ✅ 상태 emit 함수
def emit_status():
    print("[DEBUG] emit_status() 실행 시작")
    prev_events = {
        "P1-A": None,
        "P1-B": None,
        "P2-A": None,
        "P2-B": None
    }

    while True:
        for bucket, label in [
            ("P1-A_status", "P1-A"),
            ("P1-B_status", "P1-B"),
            ("P2-A_status", "P2-A"),
            ("P2-B_status", "P2-B")
        ]:
            events = get_recent_status(bucket)
            if events:
                latest = events[0]
                if latest != prev_events[label]:
                    print(f"[Influx] {label} 상태 변경: {latest}")
                    socketio.emit('status_update', {
                        label: {'event_type': latest}
                    })
                    prev_events[label] = latest
        socketio.sleep(1)

# ✅ 메인 페이지 라우팅 추가
def index():
    return render_template("index.html")
app.route("/")(index)

# ✅ 유용성 페이지 라우팅 추가
def usefulness():
    return render_template("usefulness.html")
app.route("/usefulness")(usefulness)

# ✅ 생산 추이 데이터 API 추가
def floor_to_hour(dt):
    return dt.replace(minute=0, second=0, microsecond=0)

@app.route("/get_production_data", methods=["POST"])
def get_production_data():
    try:
        # 🔍 실제 생산 완료만 카운트 (status == "finish")
        query = '''
        from(bucket: "process")
        |> range(start: -12h)
        |> filter(fn: (r) => r._measurement == "process_log" and r._field == "status" and r._value == "finish")
        |> group(columns:["line_id"])
        |> aggregateWindow(every: 1h, fn: count, createEmpty: false)
        '''
        result = influx_client.query_api().query(org=INFLUX_ORG, query=query)

        # 기준 시간대 정의 (09:00 ~ 18:00)
        time_slots = [f"{h:02}:00" for h in range(9, 19)]
        label_set = set(time_slots)
        data_by_line = {line: {} for line in ["P1-A", "P1-B", "P2-A", "P2-B"]}

        for table in result:
            for record in table.records:
                time_label = record.get_time().strftime("%H:%M")
                line_id = record.values.get("line_id")
                value = int(record.get_value())
                if line_id in data_by_line and time_label in label_set:
                    data_by_line[line_id][time_label] = value

        response = {"labels": time_slots}
        for line_id in data_by_line:
            response[line_id] = [data_by_line[line_id].get(t, 0) for t in time_slots]

        return jsonify(response)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ✅ 챗봇/Langgraph 설정
class State(TypedDict):
    messages: Annotated[list, add_messages]

# ✅ InfluxDB 쿼리용 Tool 함수
def influxdb_flux_query_tool(process_id: str):
    query_api = influx_client.query_api()

    if not process_id or not isinstance(process_id, str):
        return "올바른 process_id 또는 line_id를 입력해주세요 (예: 'P1' 또는 'P1-A')"

    if "-" in process_id:
        query = f'''
        from(bucket: "process")
        |> range(start: -1h)
        |> filter(fn: (r) => r._measurement == "process_log")
        |> filter(fn: (r) => r.line_id == "{process_id}")
        '''
    else:
        query = f'''
        from(bucket: "process")
        |> range(start: -1h)
        |> filter(fn: (r) => r._measurement == "process_log")
        |> filter(fn: (r) => r.process_id == "{process_id}")
        '''

    try:
        tables = query_api.query(query)
        logs = []
        for table in tables:
            for record in table.records:
                logs.append(f"{record.get_time()}: {record.get_value()}")
        return "\n".join(logs) if logs else f"{process_id}에 대한 로그가 없습니다."
    except Exception as e:
        return f"로그 조회 중 오류 발생: {e}"

# ✅ LangGraph 챗봇 생성 함수
def create_langgraph_chatbot():
    tools = [Tool(
        name="query_process_logs",
        func=influxdb_flux_query_tool,
        description="Input은 'P1' 또는 'P1-A'와 같은 process_id나 line_id입니다."
    )]

    system_prompt = (
        "너는 제조 공정 데이터를 분석해주는 전문 챗봇이야.\n"
        "📌 사용자의 질문에 대해 반드시 단 하나의 명확한 문장으로 요약해서 답변해줘.\n"
        "예: '최근 1시간 동안 P2-A의 평균 가동률은 74.3%입니다.'\n"
        "\n"
        "✅ 반드시 다음 기준을 따라야 해:\n"
        "- 계산 과정을 설명하거나 로그 데이터를 나열하지 마.\n"
        "- 수식, 표, 코드 블록은 절대 포함하지 마.\n"
        "- '~다음과 같습니다:', '~입니다:', '~아래와 같습니다:' 같은 패턴은 사용 금지.\n"
        "- 답변 문장은 반드시 마침표로 끝내고, 말줄임표(...) 사용하지 마.\n"
        "- '-'로 나열하는 경우는 꼭 줄바꿈해서 표시해줘.\n"
        "- 단위(%, 회 등)는 생략하지 말고 정확히 표시해.\n"
        "- 의미 없는 중복 표현은 줄이고, 최대한 간결하게 말해.\n"
        "\n"
        "⛔ 아래는 금지 예시야:\n"
        "- 다음과 같습니다: ... (금지)\n"
        "- 아래 내용을 참고하세요. (금지)\n"
        "- 90%의 가동률입니다. 감사합니다. (감사는 금지)\n"
        "\n"
        "🧠 네 역할은 '데이터를 해석해서 정리된 문장으로 전달'하는 것이지, 분석 보고서가 아니야.\n"
        "마지막으로, 반드시 문장이 하나라도 '정확한 수치'를 포함해야 해."
    )

    llm = ChatOpenAI(model="gpt-4o", temperature=0).with_config(
        configurable={"system_message": system_prompt}
    ).bind_tools(tools)

    def chatbot_node(state: State):
        return {"messages": [llm.invoke(state["messages"])]}

    graph = StateGraph(State)
    graph.add_node("chatbot", chatbot_node)
    graph.add_node("tools", ToolNode(tools=tools))
    graph.add_conditional_edges("chatbot", tools_condition)
    graph.add_edge("tools", "chatbot")
    graph.set_entry_point("chatbot")
    graph.set_finish_point("chatbot")
    return graph.compile(checkpointer=MemorySaver())

# ✅ /chat 라우팅: LangGraph 기반 챗봇
@app.route("/chat", methods=["POST"])
def chat():
    user_message = request.json.get("message", "")
    try:
        graph = create_langgraph_chatbot()
        config = RunnableConfig(recursion_limit=10, configurable={"thread_id": "web-user"})
        response_text = ""
        for event in graph.stream({"messages": [{"role": "user", "content": user_message}]}, config=config):
            for value in event.values():
                if "messages" in value and value["messages"]:
                    response_text = value["messages"][-1].content
        return jsonify({"reply": response_text})
    except Exception as e:
        return jsonify({"reply": f"❌ LangGraph 챗봇 오류: {str(e)}"}), 500


# ✅ 최근 이벤트 상태 조회 함수
def get_recent_status(bucket):
    query = f'''
    from(bucket: "{bucket}")
      |> range(start: -30s)
      |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_type")
      |> sort(columns: ["_time"], desc: true)
      |> limit(n: 3)
    '''
    result = influx_client.query_api().query(org=INFLUX_ORG, query=query)

    events = []
    for table in result:
        for record in table.records:
            events.append(record.get_value())
    return events if events else None

# ✅ 클라이언트 최초 연결 시 상태 전송
@socketio.on('connect')
def handle_connect():
    print('Client connected')
    for bucket, label in [
        ("P1-A_status", "P1-A"),
        ("P1-B_status", "P1-B"),
        ("P2-A_status", "P2-A"),
        ("P2-B_status", "P2-B")
    ]:
        events = get_recent_status(bucket)
        if events:
            latest = events[0]
            socketio.emit('status_update', {
                label: {'event_type': latest}
            })


# ✅ 보고서 페이지
@app.route("/report")
def report_page():
    return render_template("report.html")

# ✅ 보고서 생성 API
@app.route("/generate_report", methods=["POST"])
def generate_report():
    data = request.json
    process = data.get("process")
    range_str = data.get("range")
    query = f'''
    from(bucket: "{process}_status")
      |> range(start: -{range_str})
      |> filter(fn: (r) => r._measurement == "status_log")
      |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
      |> keep(columns: ["_time", "available", "event_type"])
    '''
    tables = influx_client.query_api().query(query)
    total = 0
    available_sum = 0
    failure_count = 0
    time_labels = []
    available_values = []
    failure_values = []
    for table in tables:
        for record in table.records:
            available = record.values.get("available", 0)
            event_type = record.values.get("event_type", "")
            timestamp = record.values["_time"].strftime("%H:%M")
            total += 1
            available_sum += available
            if event_type == "failure":
                failure_count += 1
            time_labels.append(timestamp)
            available_values.append(round(available, 2))
            failure_values.append(1 if event_type == "failure" else 0)
    avg_avail = round((available_sum / total) * 100, 1) if total else 0
    prompt = f"""
공정명: {process}
기간: 최근 {range_str}
가동률 평균: {avg_avail}%
고장 횟수: {failure_count}회

위 데이터를 바탕으로 제조 공정 보고서를 작성해줘. 다음 항목을 포함해줘:
1. 공정 요약
2. 주요 이슈
3. 대응 조치
4. 향후 제언
"""
    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 제조공정 보고서를 작성하는 AI 비서야."},
                {"role": "user", "content": prompt}
            ]
        )
        return jsonify({
            "report": response.choices[0].message.content,
            "labels": time_labels,
            "available": available_values,
            "failures": failure_values
        })
    except Exception as e:
        print(f"오류 발생: {e}")
        return jsonify({"error": str(e)}), 500

# ✅ 보고서 다운로드 API
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    try:
        data = request.json
        text = data.get("report", "")
        avail_img_b64 = data.get("availabilityImage", "")
        fail_img_b64 = data.get("failureImage", "")
        doc = Document()
        doc.add_heading("\ud83d\udcc4 스마트 제조 보고서", 0)
        doc.add_paragraph(text)
        def add_image(doc, b64_string, title):
            if b64_string and "base64," in b64_string:
                try:
                    doc.add_paragraph(title)
                    image_data = base64.b64decode(b64_string.split(",")[-1])
                    image_stream = BytesIO(image_data)
                    doc.add_picture(image_stream, width=Inches(2.75))
                except Exception as img_err:
                    print(f"이미지 디코딩 오류: {img_err}")
        add_image(doc, avail_img_b64, "\u2705 가동률 변화")
        add_image(doc, fail_img_b64, "\ud83d\udcca 고장 발생 분포")
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name="제조_보고서.docx")
    except Exception as e:
        print(f"\ud83d\udd1b DOCX 생성 오류: {e}")
        return jsonify({"error": "파일 생성 실패"}), 500


# ✅ 서버 실행
if __name__ == "__main__":
    socketio.start_background_task(target=emit_status)
    socketio.run(app, host='0.0.0.0', port=5000, debug=False, use_reloader=False, log_output=True)