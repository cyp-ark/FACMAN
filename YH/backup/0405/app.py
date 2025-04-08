import eventlet
eventlet.monkey_patch()

import os
import base64
from io import BytesIO
from flask import Flask, render_template, request, jsonify, send_file
from flask_socketio import SocketIO
from influxdb_client import InfluxDBClient
from dotenv import load_dotenv
from flask_cors import CORS
from docx import Document
from docx.shared import Inches

# ✅ 환경 변수 로드
load_dotenv()

# ✅ Flask 앱 및 SocketIO 초기화
app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode="eventlet")

# ✅ InfluxDB 클라이언트 설정
INFLUX_URL = os.getenv("INFLUX_URL")
INFLUX_TOKEN = os.getenv("INFLUX_TOKEN")
INFLUX_ORG = os.getenv("INFLUX_ORG")
influx_client = InfluxDBClient(url=INFLUX_URL, token=INFLUX_TOKEN, org=INFLUX_ORG)

# ✅ 상태 emit 함수
def emit_status():
    print("[DEBUG] emit_status() 실행 시작")
    prev_events = {
        "P1-A": None,
        "P1-B": None,
        "P2-A": None,
        "P2-B": None
    }

    while True:
        for bucket, label in [
            ("P1-A_status", "P1-A"),
            ("P1-B_status", "P1-B"),
            ("P2-A_status", "P2-A"),
            ("P2-B_status", "P2-B")
        ]:
            events = get_recent_status(bucket)
            if events:
                latest = events[0]
                if latest != prev_events[label]:
                    print(f"[Influx] {label} 상태 변경: {latest}")
                    socketio.emit('status_update', {
                        label: {'event_type': latest}
                    })
                    prev_events[label] = latest
        socketio.sleep(1)


# ✅ 메인 페이지 라우팅 추가
def index():
    return render_template("index.html")
app.route("/")(index)

# ✅ 유용성 페이지 라우팅 추가
def usefulness():
    return render_template("usefulness.html")
app.route("/usefulness")(usefulness)

# ✨ ✅ 유용성 데이터 API 라우팅 추가
@app.route("/get_usefulness_data", methods=["POST"])
def get_usefulness_data():
    try:
        data = request.get_json()
        process = data.get("process", "P1-A")
        time_range = data.get("range", "6h")

        # 예시 Flux 코드는 평균 available 계산을 가지고 올 수 있어요
        query = f'''
        from(bucket: "{process}_status")
          |> range(start: -{time_range})
          |> filter(fn: (r) => r._measurement == "status_log" and r._field == "available")
          |> aggregateWindow(every: 1h, fn: mean, createEmpty: false)
          |> yield(name: "mean")
        '''

        result = influx_client.query_api().query(org=INFLUX_ORG, query=query)
        labels, values = [], []
        for table in result:
            for record in table.records:
                labels.append(record.get_time().strftime("%H:%M"))
                values.append(round(record.get_value() * 100, 1))  # 턀시티지

        return jsonify({
            "labels": labels,
            "values": values
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ✅ 최근 이벤트 상태 조회 함수
def get_recent_status(bucket):
    query = f'''
    from(bucket: "{bucket}")
      |> range(start: -30s)
      |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_type")
      |> sort(columns: ["_time"], desc: true)
      |> limit(n: 3)
    '''
    result = influx_client.query_api().query(org=INFLUX_ORG, query=query)

    events = []
    for table in result:
        for record in table.records:
            events.append(record.get_value())
    return events if events else None


# ✅ 클라이언트 최초 연결 시 상태 전송
@socketio.on('connect')
def handle_connect():
    print('Client connected')
    for bucket, label in [
        ("P1-A_status", "P1-A"),
        ("P1-B_status", "P1-B"),
        ("P2-A_status", "P2-A"),
        ("P2-B_status", "P2-B")
    ]:
        events = get_recent_status(bucket)
        if events:
            latest = events[0]
            socketio.emit('status_update', {
                label: {'event_type': latest}
            })

# ✅ 보고서 페이지
@app.route("/report")
def report_page():
    return render_template("report.html")

# ✅ 보고서 생성 API
@app.route("/generate_report", methods=["POST"])
def generate_report():
    data = request.json
    process = data.get("process")
    range_str = data.get("range")
    query = f'''
    from(bucket: "{process}_status")
      |> range(start: -{range_str})
      |> filter(fn: (r) => r._measurement == "status_log")
      |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
      |> keep(columns: ["_time", "available", "event_type"])
    '''
    tables = influx_client.query_api().query(query)
    total = 0
    available_sum = 0
    failure_count = 0
    time_labels = []
    available_values = []
    failure_values = []
    for table in tables:
        for record in table.records:
            available = record.values.get("available", 0)
            event_type = record.values.get("event_type", "")
            timestamp = record.values["_time"].strftime("%H:%M")
            total += 1
            available_sum += available
            if event_type == "failure":
                failure_count += 1
            time_labels.append(timestamp)
            available_values.append(round(available, 2))
            failure_values.append(1 if event_type == "failure" else 0)
    avg_avail = round((available_sum / total) * 100, 1) if total else 0
    prompt = f"""
공정명: {process}
기간: 최근 {range_str}
가동률 평균: {avg_avail}%
고장 횟수: {failure_count}회

위 데이터를 바탕으로 제조 공정 보고서를 작성해줘. 다음 항목을 포함해줘:
1. 공정 요약
2. 주요 이슈
3. 대응 조치
4. 향후 제언
"""
    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 제조공정 보고서를 작성하는 AI 비서야."},
                {"role": "user", "content": prompt}
            ]
        )
        return jsonify({
            "report": response.choices[0].message.content,
            "labels": time_labels,
            "available": available_values,
            "failures": failure_values
        })
    except Exception as e:
        print(f"오류 발생: {e}")
        return jsonify({"error": str(e)}), 500

# ✅ 보고서 다운로드 API
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    try:
        data = request.json
        text = data.get("report", "")
        avail_img_b64 = data.get("availabilityImage", "")
        fail_img_b64 = data.get("failureImage", "")
        doc = Document()
        doc.add_heading("\ud83d\udcc4 스마트 제조 보고서", 0)
        doc.add_paragraph(text)
        def add_image(doc, b64_string, title):
            if b64_string and "base64," in b64_string:
                try:
                    doc.add_paragraph(title)
                    image_data = base64.b64decode(b64_string.split(",")[-1])
                    image_stream = BytesIO(image_data)
                    doc.add_picture(image_stream, width=Inches(2.75))
                except Exception as img_err:
                    print(f"이미지 디코딩 오류: {img_err}")
        add_image(doc, avail_img_b64, "\u2705 가동률 변화")
        add_image(doc, fail_img_b64, "\ud83d\udcca 고장 발생 분포")
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name="제조_보고서.docx")
    except Exception as e:
        print(f"\ud83d\udd1b DOCX 생성 오류: {e}")
        return jsonify({"error": "파일 생성 실패"}), 500

@app.route("/chat", methods=["POST"])
def chat():
    user_message = request.json.get("message", "")
    intent_prompt = f"""
너는 AI 에이전트야. 사용자의 질문을 보고, 아래 중 하나로 분류해줘.
- influx: 제조 공정이나 설비 상태, 고장 등 InfluxDB에서 데이터를 가져와야 하는 질문
- web: 외부 뉴스, 시세, 일반 정보 등 웹 검색이 필요한 질문
- gpt: 일반적인 지식이나 개념 설명, 잡담 등

질문: "{user_message}"
위 질문은 어떤 유형이야? 반드시 influx / web / gpt 중 하나만 말해줘.
"""
    try:
        intent_res = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 사용자의 질문 intent를 분석하는 판단 에이전트야."},
                {"role": "user", "content": intent_prompt}
            ]
        )
        intent_raw = intent_res.choices[0].message.content.strip().lower()
        intent = intent_raw if intent_raw in ["influx", "web", "gpt"] else "gpt"
        if intent == "influx":
            reply = handle_influx_query(user_message)
        elif intent == "web":
            reply = "\ud83d\udd0d 웹 검색 기능은 현재 준비 중입니다."
        else:
            reply = handle_gpt_query(user_message)
        return jsonify({"reply": reply})
    except Exception as e:
        return jsonify({"reply": f"\u26a0\ufe0f 오류 발생: {str(e)}"}), 500

# ✅ GPT 처리 함수
def handle_gpt_query(user_message):
    gpt_reply = openai_client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "너는 친절한 제조 공정 AI 비서야."},
            {"role": "user", "content": user_message}
        ]
    )
    return gpt_reply.choices[0].message.content.strip()

# ✅ Influx 처리 함수 (GPT 기반 쿼리 생성 + 실행)
def handle_influx_query(user_message):
    query_prompt = f"""
너는 InfluxDB 전문가야. 다음과 같이 정확하게 Flux 쿼리를 생성해줘.

⚠️ 주의사항:
- 버킷 이름은 반드시 아래 중 하나로 고정해야 해:
  - "P1-A_status"
  - "P1-B_status"
  - "P2-A_status"
  - "P2-B_status"
- "your_bucket" 같은 표현은 절대 사용하면 안 돼.
- 쿼리는 실행 가능한 형태여야 하고, 결과에 _value 또는 available, event_type 필드가 있어야 해.

사용자 질문: "{user_message}"
Flux 쿼리만 반환해줘. 설명은 필요 없어.
"""
    try:
        gpt_query_res = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 InfluxDB 쿼리 생성 전문가야."},
                {"role": "user", "content": query_prompt}
            ]
        )
        flux_query = gpt_query_res.choices[0].message.content.strip()
        print("✅ 생성된 쿼리:\n", flux_query)
        result_tables = influx_client.query_api().query(flux_query)
        result_rows = []
        for table in result_tables:
            for record in table.records:
                values = record.values
                time_str = str(values.get("_time", "(시간 없음)"))
                field_str = str(values.get("_field", "(필드 없음)"))
                value_str = str(values.get("_value", values.get("value", "(값 없음)")))
                result_rows.append(f"{time_str} - {field_str} = {value_str}")
        if not result_rows:
            return "📍 InfluxDB에서 결과를 찾을 수 없습니다."
        return "📊 InfluxDB 응답 결과:\n" + "\n".join(result_rows[:10])
    except Exception as e:
        return f"⚠️ 쿼리 처리 중 오류 발생: {str(e)}"

# ✅ 서버 실행
if __name__ == "__main__":
    socketio.start_background_task(target=emit_status)
    socketio.run(app, host='0.0.0.0', port=5000, debug=False, use_reloader=False, log_output=True)