from flask import Flask, request, jsonify, render_template, send_file
from flask_socketio import SocketIO, emit
from influxdb_client import InfluxDBClient
from docx import Document
from docx.shared import Inches
from io import BytesIO
from openai import OpenAI
import base64
from threading import Thread
from time import sleep
import os
from dotenv import load_dotenv

# ✅ 환경 변수 로드
load_dotenv()

# ✅ Flask 앱 초기화
app = Flask(__name__)
socketio = SocketIO(app)  # ✅ SocketIO 활성화

# ✅ .env에서 민감 정보 불러오기
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
INFLUX_URL = os.getenv("INFLUX_URL")
INFLUX_TOKEN = os.getenv("INFLUX_TOKEN")
INFLUX_ORG = os.getenv("INFLUX_ORG")

influx_client = InfluxDBClient(url=INFLUX_URL, token=INFLUX_TOKEN, org=INFLUX_ORG)

# ✅ 메인 대시보드 페이지
@app.route("/")
def index():
    return render_template("index.html")

# ✅ 공정 상태를 주기적으로 감지하고 WebSocket으로 전송
def emit_status_loop():
    while True:
        try:
            statuses = {}

            for process in ["P1", "P2"]:
                try:
                    query = f'''
                    from(bucket: "{process}_status")
                      |> range(start: -5m)
                      |> filter(fn: (r) => r._measurement == "status_log" and r._field == "available")
                      |> last()
                    '''

                    print(f"🔍 {process} 쿼리 실행 중...")

                    tables = influx_client.query_api().query(query)
                    found = False

                    for table in tables:
                        for record in table.records:
                            value = record.get_value()
                            print(f"✅ {process} 상태값:", value)
                            statuses[process] = int(value)
                            found = True

                    if not found:
                        print(f"⚠️ {process}: 최근 5분 내 'available' 데이터 없음")

                except Exception as e:
                    print(f"❌ {process} 쿼리 오류:", e)

            print("📤 emit할 상태:", statuses)
            socketio.emit('status_update', statuses)

        except Exception as e:
            print("🔥 상태 송신 오류:", e)

        sleep(5)


# ✅ 백그라운드 스레드 실행
@socketio.on('connect')
def handle_connect():
    print("📡 클라이언트 WebSocket 연결됨")


# ✅ 보고서 페이지
@app.route("/report")
def report_page():
    return render_template("report.html")

# ✅ 보고서 생성 API
@app.route("/generate_report", methods=["POST"])
def generate_report():
    data = request.json
    process = data.get("process")
    range_str = data.get("range")

    query = f'''
    from(bucket: "{process}_status")
      |> range(start: -{range_str})
      |> filter(fn: (r) => r._measurement == "status_log")
      |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
      |> keep(columns: ["_time", "available", "event_type"])
    '''

    tables = influx_client.query_api().query(query)

    total = 0
    available_sum = 0
    failure_count = 0
    time_labels = []
    available_values = []
    failure_values = []

    for table in tables:
        for record in table.records:
            available = record.values.get("available", 0)
            event_type = record.values.get("event_type", "")
            timestamp = record.values["_time"].strftime("%H:%M")

            total += 1
            available_sum += available
            if event_type == "failure":
                failure_count += 1

            time_labels.append(timestamp)
            available_values.append(round(available, 2))
            failure_values.append(1 if event_type == "failure" else 0)

    avg_avail = round((available_sum / total) * 100, 1) if total else 0

    prompt = f"""
공정명: {process}
기간: 최근 {range_str}
가동률 평균: {avg_avail}%
고장 횟수: {failure_count}회

위 데이터를 바탕으로 제조 공정 보고서를 작성해줘. 다음 항목을 포함해줘:
1. 공정 요약
2. 주요 이슈
3. 대응 조치
4. 향후 제언
"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 제조공정 보고서를 작성하는 AI 비서야."},
                {"role": "user", "content": prompt}
            ]
        )
        return jsonify({
            "report": response.choices[0].message.content,
            "labels": time_labels,
            "available": available_values,
            "failures": failure_values
        })
    except Exception as e:
        print(f"오류 발생: {e}")
        return jsonify({"error": str(e)}), 500

# ✅ 보고서 다운로드 API
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    try:
        data = request.json
        text = data.get("report", "")
        avail_img_b64 = data.get("availabilityImage", "")
        fail_img_b64 = data.get("failureImage", "")

        doc = Document()
        doc.add_heading("📄 스마트 제조 보고서", 0)
        doc.add_paragraph(text)

        def add_image(doc, b64_string, title):
            if b64_string and "base64," in b64_string:
                try:
                    doc.add_paragraph(title)
                    image_data = base64.b64decode(b64_string.split(",")[-1])
                    image_stream = BytesIO(image_data)
                    doc.add_picture(image_stream, width=Inches(2.75))
                except Exception as img_err:
                    print(f"이미지 디코딩 오류: {img_err}")

        add_image(doc, avail_img_b64, "✅ 가동률 변화")
        add_image(doc, fail_img_b64, "📊 고장 발생 분포")

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name="제조_보고서.docx")

    except Exception as e:
        print(f"📛 DOCX 생성 오류: {e}")
        return jsonify({"error": "파일 생성 실패"}), 500


@app.route("/chat", methods=["POST"])
def chat():
    user_message = request.json.get("message", "")

    # ✅ 1단계: intent 분류 프롬프트
    intent_prompt = f"""
너는 AI 에이전트야. 사용자의 질문을 보고, 아래 중 하나로 분류해줘.
- influx: 제조 공정이나 설비 상태, 고장 등 InfluxDB에서 데이터를 가져와야 하는 질문
- web: 외부 뉴스, 시세, 일반 정보 등 웹 검색이 필요한 질문
- gpt: 일반적인 지식이나 개념 설명, 잡담 등

질문: "{user_message}"
위 질문은 어떤 유형이야? 반드시 influx / web / gpt 중 하나만 말해줘.
"""

    try:
        intent_res = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 사용자의 질문 intent를 분석하는 판단 에이전트야."},
                {"role": "user", "content": intent_prompt}
            ]
        )
        intent_raw = intent_res.choices[0].message.content.strip().lower()
        intent = intent_raw if intent_raw in ["influx", "web", "gpt"] else "gpt"

        # ✅ 2단계: intent에 따라 처리
        if intent == "influx":
            reply = handle_influx_query(user_message)
        elif intent == "web":
            reply = "🔍 웹 검색 기능은 현재 준비 중입니다."
        else:
            reply = handle_gpt_query(user_message)

        return jsonify({"reply": reply})
    except Exception as e:
        return jsonify({"reply": f"⚠️ 오류 발생: {str(e)}"}), 500


# ✅ GPT 처리 함수
def handle_gpt_query(user_message):
    gpt_reply = openai_client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "너는 친절한 제조 공정 AI 비서야."},
            {"role": "user", "content": user_message}
        ]
    )
    return gpt_reply.choices[0].message.content.strip()


# ✅ Influx 처리 함수 (GPT 기반 쿼리 생성 + 실행)
def handle_influx_query(user_message):
    # Step 1: 사용자 질문을 Flux 쿼리로 변환 (버킷명 고정 안내 포함)
    query_prompt = f"""
너는 InfluxDB 전문가야. 다음과 같이 정확하게 Flux 쿼리를 생성해줘.

⚠️ 주의사항:
- 버킷 이름은 반드시 아래 중 하나로 고정해야 해:
  - "P1_status"
  - "P2_status"
- "your_bucket" 같은 표현은 절대 사용하면 안 돼.
- 쿼리는 실행 가능한 형태여야 하고, 결과에 _value 또는 available, event_type 필드가 있어야 해.

사용자 질문: "{user_message}"
Flux 쿼리만 반환해줘. 설명은 필요 없어.
"""

    try:
        gpt_query_res = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 InfluxDB 쿼리 생성 전문가야."},
                {"role": "user", "content": query_prompt}
            ]
        )
        flux_query = gpt_query_res.choices[0].message.content.strip()
        print("✅ 생성된 쿼리:\n", flux_query)

        # Step 2: InfluxDB 쿼리 실행
        result_tables = influx_client.query_api().query(flux_query)

        # Step 3: 결과 파싱
        result_rows = []
        for table in result_tables:
            for record in table.records:
                values = record.values
                time_str = str(values.get("_time", "(시간 없음)"))
                field_str = str(values.get("_field", "(필드 없음)"))
                value_str = str(values.get("_value", values.get("value", "(값 없음)")))

                result_rows.append(f"{time_str} - {field_str} = {value_str}")

        if not result_rows:
            return "📭 InfluxDB에서 결과를 찾을 수 없습니다."

        return "📊 InfluxDB 응답 결과:\n" + "\n".join(result_rows[:10])

    except Exception as e:
        return f"⚠️ 쿼리 처리 중 오류 발생: {str(e)}"


# ✅ 앱 실행
if __name__ == "__main__":
    Thread(target=emit_status_loop, daemon=True).start()  # ✅ 상태 감지 스레드 시작
    socketio.run(app, debug=True, host="0.0.0.0", port=5000)
